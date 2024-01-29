from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from collections import OrderedDict
import sys
import os

A4W, A4H = 8.3, 11.7
doc = Document()

comments_syntax  = {
    "sh":"#",
    "c":"//",
    "js":"//",
    "java":"//"
}

def get_files_inorder(source_path):
    all_files = os.listdir(source_path)

    sorted_files = dict()

    pure_files = [file for file in all_files if os.path.isfile(os.path.join(source_path, file))]

    for f in pure_files:
        with open(source_path+"/"+f,'r') as file:
            content = file.read().strip().split("\n")[0]
            if(content[0]==comments_syntax[f.split(".")[-1]]):
                content = content[1:].strip().split(".")[0].strip()
                sorted_files[content] = source_path+"/"+f


    return OrderedDict(sorted(sorted_files.items(),key=lambda x:x[0]))     


def remove_formatting(paragraph):
    for run in paragraph.runs:
        run.font.size = None
        run.font.bold = None
        run.font.italic = None
        run.font.underline = None
        run.font.color.rgb = None

def set_heading_formating(heading):
    for run in heading.runs:
        run.font.size = 25

def set_text_color_to_black(element):
    for run in element.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set color to black

def build_assignment(files,name):
    # section = doc.sections[0]
    doc.page_height = Inches(A4H)  # A4 height in points
    doc.page_width = Inches(A4W)   # A4 width in points
    head = doc.add_heading('Assignment', level=0)
    head.style = 'Normal'
    set_heading_formating(head)
    # paragraph = doc.add_paragraph('This is a paragraph.')

    for val in files.values():
        with open(val,'r') as file:
            content = file.read().strip()
            line = content.split("\n")
            heading = line[0]
            heading = heading[1:].strip()
            code = "\n".join(line[1:])

            head = doc.add_heading(heading,level=2)
            para = doc.add_paragraph(code)
            para.paragraph_format.left_indent = Inches(0.3)
            set_text_color_to_black(head)

    doc.save(name+".docx")


if __name__ == "__main__":

    try:
        source_path = sys.argv[1]
        doc_name = sys.argv[2]

        files = (get_files_inorder(source_path))
        build_assignment(files,doc_name.strip())
        # build_assignment(doc_name)

    except IndexError as e:
        print("Source or doc name not provided")
    